import csv
import glob
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import pycountry

Shipping_Data_Folder = "All Rewards - Jul 1 11pm - usps"
Shipping_Details_From_Kickstarter_Style_Export_folder="KickstarterStyleReport Jul 1"
LIST_FOCUS = "INTERNATIONAL"   	#Options: DOMESTIC, INTERNATIONAL, ALL
							#DOMESTIC will NOT show non-subscribed people
							
DOCUMENT_TYPE = "PACKING_SLIP"   #Options: PACKING_SLIP, FRONT_POSTAGE
							
Main_Folder=r"D:\Programs\Hubic\Maker\Business\FPGA\Kickstarter Delivery\Rewards Reports"

files = glob.glob(Main_Folder+'\\'+Shipping_Data_Folder+'\\*.csv')
shipping_details_files = glob.glob(Main_Folder+'\\'+Shipping_Details_From_Kickstarter_Style_Export_folder+'\\*.csv')

def openCSV(file):
	csvfile = open(file, 'rt',encoding="utf8")
	csv_reader = csv.reader(csvfile, delimiter=',', quotechar='"')
	return csv_reader

def makePage(document, row):
	First=row[0]
	Last=row[2]
	Address1=row[4]
	Address2=row[5]
	Address3=row[6]
	City=row[7]
	State=row[8]
	Zip=row[9]
	Country_code=row[10]
	Quantity=row[17]
	Description=row[18]
	if Country_code:
		Country_name= pycountry.countries.get(alpha_2=Country_code).name
	else:
		Country_name=''
	Phone=row[12]
	Email=row[14]
	ReferenceNumber=row[15]	
	DeliveryDetails = shipping_details_dict[str(ReferenceNumber)]
	
	document.add_picture('D:\\Programs\\Hubic\\Maker\\Assets\\Logo\\Logo Crisp.png', width=Inches(5.0))	
	last_paragraph = document.paragraphs[-1] 
	last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	
	heading="Order #"+ReferenceNumber
	document.add_heading(heading, 0)
	document.add_heading(First+" "+Last, level=1)
	
	if Address2:
		if Address3:
			Address=Address1+"\n"+Address2+"\n"+Address3+"\n"
		else:
			Address=Address1+"\n"+Address2
	else:
		Address=Address1
		
	
		
	CityStateZip=City+", "+State+" "+Zip
	
	document.add_paragraph(Address+'\n'+CityStateZip)
	document.add_paragraph(Country_name)
	if Phone:
		document.add_heading("Phone number:   "+Phone, level=3)
	if DeliveryDetails:
		document.add_heading("Notes for Shipping:     "+DeliveryDetails, level=3)
	
	#document.add_paragraph('Fipsy FPGA\n\n\n', style='IntenseQuote')

	#document.add_picture('monty-truth.png', width=Inches(1.25))

	table = document.add_table(rows=1, cols=1)
	table.style = 'LightShading-Accent1'
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Your Order'
	cells = table.add_row().cells
	quantityText=""+Quantity+" x Fipsy FPGA\n\nDescription:\n"+Description+"\n"
	cells[0].text=quantityText
	
	run = cells[0].paragraphs[0].runs[0]
	font=run.font
	font.color.rgb = RGBColor(0x00, 0x00, 0x00)	
	
	document.add_heading("Instructions", level=2)
	
	Instructions="""Start by learning how to program the Fipsy FPGA here: https://www.mocomakers.com/fipsy-fpga/programmers/\n
Then learn how to code for FPGAs here: https://www.mocomakers.com/wiki/learning/"""
	paragraph = document.add_paragraph(Instructions)
	Instructions="""\n\nLastly, share your projects with our community of learners here: https://www.mocomakers.com/projects/sharing/ \n
Share your project and be entered to win $200 USD (see the link).\n"""
	paragraph.add_run(Instructions).bold = True
	
	heading="MoCo Makers"
	document.add_heading(heading, 0)
	
	last_paragraph = document.paragraphs[-1] 
	last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	
	heading="www.MoCoMakers.com"
	document.add_heading(heading, 3)
	
	last_paragraph = document.paragraphs[-1] 
	last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	

	
	"""
	cells = table.add_row().cells
	if not qty_brown:
		qty_brown=0
	if not qty_black:
		qty_black=0
	if not qty_spread:
		qty_spread=0
	total=(int(qty_brown)*5.50)+(int(qty_black)*6.75)+(int(qty_spread)*4.0)
	print "using stated: "+str(stated_total)+" for order: "+str(order_num)
	stated_total=float(stated_total)
	
	if total!=stated_total:
		print "Brown used "+str(qty_brown)
		print "Black used "+str(qty_black)
		print "qty_spread "+str(qty_spread)
		raise Exception("Total Price Mismatch: "+str(total)+" "+str(stated_total))
	cells[0].text = "TOTAL: "+'${:,.2f}'.format(total)
	
	p = document.add_paragraph('')
	p.add_run('-------Delivery Directions-------').bold
	
	document.add_paragraph(deliv_instruct)
	
	if int(qty_spread)>0:
		p = document.add_paragraph('')
		p.add_run('-------Spreading Instructions-------').bold
		
		print "Now using string:'"+str(spread_instruct)+"'"
		document.add_paragraph(spread_instruct)
	"""

	document.add_page_break()
	if Address1:
		return (document, True)
	else:
		return (document, False, ReferenceNumber, Email)

def makeLabel(row):
	First=row[0]
	Last=row[2]
	Address1=row[4]
	Address2=row[5]
	Address3=row[6]
	City=row[7]
	State=row[8]
	Zip=row[9]
	Country_code=row[10]
	Quantity=row[17]
	Description=row[18]
	if Country_code:
		Country_name= pycountry.countries.get(alpha_2=Country_code).name
	else:
		Country_name=''
	Phone=row[12]
	Email=row[14]
	ReferenceNumber=row[15]	
	
	output_str=First+" "+Last+"\n"+Address1
	
	if Address2:
		output_str=output_str+"\n"+Address2
		
	if Address3:
		output_str=output_str+"\n"+Address3
	output_str=output_str+"\n"+City+", "+State+" "+Zip
	
	if not Address2:
		output_str=output_str+"\n"
	
	if not Address3:
		output_str=output_str+"\n"
	
	return output_str
	
if __name__=="__main__":
	shipping_details_dict = {}
	for shipping_details_csv in shipping_details_files:
		if "No reward" in shipping_details_csv:
			pass
		else:
			shipping_csv_reader=openCSV(shipping_details_csv)
			is_first=True
			
			for row in shipping_csv_reader:
				if is_first:
					is_first=False
				else:
					print(str(row))
					backer_number = row[0]
					shipping_details = row[24]
					print("Backer Number: "+str(backer_number))
					print("Description: "+str(shipping_details))
					shipping_details_dict[backer_number]=shipping_details
	
	for file in files:
		if "No reward - " in file:
			pass
		else:
			csv_reader=openCSV(file)
					
			is_first=True
			
			customers=[]
			
			for row in csv_reader:
				if is_first:
					is_first=False
				else:
					#print str(row)
					customers.append(row)
			document = Document()
			save_file = Main_Folder+"\\Outputs\\"+file.split('\\')[-1].split('.csv')[0]+" - Output.docx"
			
			Description = file.split('\\')[-1].split(' - ')[0].split('USD')[1].strip()
			if Description=="One Fipsy FPGA" or Description=="One Pre-soldered Fipsy FPGA":
				quantity=1
			elif Description=="Two Fipsy Pack" or Description=="Pre-soldered Fipsy FPGA Two-Pack":
				quantity=2
			elif Description=="Fipsy FPGA Four-Pack":
				quantity=4
			elif Description=="Fipsy FPGA Ten-Pack" or Description=="Pre-soldered FPGA Ten-Pack":
				quantity=10
			else:
				raise Exception("Failed to match quantity to file name")
			
			quantity=str(quantity)
			
			print(str(file))
			output_str=""  #For Lables
			
			for row in customers:
				Country_code=row[10]
				
				if LIST_FOCUS=="DOMESTIC" and Country_code!='US':
					pass
				elif LIST_FOCUS=="INTERNATIONAL" and Country_code=='US':
					pass
				else:			
					row.append(quantity)
					row.append(Description)
					if DOCUMENT_TYPE=="PACKING_SLIP":
						Email = row[14]
						Address1=row[4]
						
						if Address1:
							results = makePage(document, row)
							document = results[0]
							is_valid = results[1]
							if not is_valid:
								raise Exception("Failed for: "+"email: "+results[3]+" "+"Order Reference:"+results[2])
						else:
							print("Failed to find address for: "+Email)
						
					else:
						if output_str:
							output_str = output_str +"\n"+ makeLabel(row)
						else:
							output_str = output_str + makeLabel(row)
				
			if DOCUMENT_TYPE=="PACKING_SLIP":
				document.save(save_file)
			else:
				print(output_str)
				print("Done creating labels..")
				with open('labels - '+Description+'.txt', 'w') as file:
					file.write(output_str)
			
	print("Now Done")